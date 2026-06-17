"""Render markdown to DOCX via python-docx. Supports headings, paragraphs,
inline emphasis, lists, code blocks, tables, hr."""
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from markdown_it import MarkdownIt


_LEADING_H1 = re.compile(r"\A\s*#\s+[^\n]+\n+")


def _add_inline(paragraph, tokens) -> None:
    """Walk markdown-it inline tokens onto a paragraph as a series of runs
    with the right bold/italic/code/link styling."""
    bold = italic = code = False
    link_href: str | None = None
    for t in tokens:
        if t.type == "text":
            run = paragraph.add_run(t.content)
            run.bold = bold
            run.italic = italic
            if code:
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            if link_href:
                run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
                run.font.underline = True
        elif t.type in ("softbreak", "hardbreak"):
            paragraph.add_run("\n")
        elif t.type == "code_inline":
            run = paragraph.add_run(t.content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif t.type == "strong_open":
            bold = True
        elif t.type == "strong_close":
            bold = False
        elif t.type == "em_open":
            italic = True
        elif t.type == "em_close":
            italic = False
        elif t.type == "link_open":
            attrs = t.attrs or []
            if isinstance(attrs, list):
                link_href = next((a[1] for a in attrs if a[0] == "href"), "")
            else:
                link_href = attrs.get("href", "")
        elif t.type == "link_close":
            link_href = None


def _add_code_block(doc: Document, content: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run(content.rstrip())
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.cell(ri, ci)
            text = row[ci] if ci < len(row) else ""
            cell.text = text
            if ri == 0:
                # Header row: bold the run.
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def _list_item_text(tokens, start: int) -> tuple[str, int]:
    """Render an inline-only list item to plain text. Returns (text, next_index)."""
    j = start
    text_parts: list[str] = []
    while j < len(tokens) and tokens[j].type != "list_item_close":
        if tokens[j].type == "inline":
            for c in tokens[j].children or []:
                if c.type == "text":
                    text_parts.append(c.content)
                elif c.type == "code_inline":
                    text_parts.append(c.content)
                elif c.type in ("softbreak", "hardbreak"):
                    text_parts.append(" ")
        j += 1
    return "".join(text_parts).strip(), j


def build_docx(path: str, title: str, body_md: str) -> None:
    body_md = _LEADING_H1.sub("", body_md or "", count=1)
    md = MarkdownIt("commonmark", {"html": False}).enable("table")
    tokens = md.parse(body_md or "")

    doc = Document()
    # Document defaults — Calibri 11 is the modern Word baseline.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(22)

    i = 0
    while i < len(tokens):
        t = tokens[i]
        if t.type == "heading_open":
            level = int(t.tag[1])
            inline = tokens[i + 1].children if i + 1 < len(tokens) and tokens[i + 1].type == "inline" else []
            heading_level = max(1, min(level, 4))
            p = doc.add_heading(level=heading_level)
            _add_inline(p, inline or [])
            i += 3
            continue
        if t.type == "paragraph_open":
            inline = tokens[i + 1].children if i + 1 < len(tokens) and tokens[i + 1].type == "inline" else []
            p = doc.add_paragraph()
            _add_inline(p, inline or [])
            i += 3
            continue
        if t.type in ("bullet_list_open", "ordered_list_open"):
            ordered = (t.type == "ordered_list_open")
            depth = 1
            j = i + 1
            counter = 1
            while j < len(tokens) and depth:
                tj = tokens[j]
                if tj.type in ("bullet_list_open", "ordered_list_open"):
                    depth += 1
                elif tj.type in ("bullet_list_close", "ordered_list_close"):
                    depth -= 1
                    if depth == 0:
                        break
                elif tj.type == "list_item_open" and depth == 1:
                    text, next_j = _list_item_text(tokens, j + 1)
                    style_name = "List Number" if ordered else "List Bullet"
                    try:
                        doc.add_paragraph(text, style=style_name)
                    except KeyError:
                        # Fall back to a manual prefix if the doc style isn't available.
                        prefix = f"{counter}. " if ordered else "• "
                        doc.add_paragraph(prefix + text)
                    counter += 1
                    j = next_j
                j += 1
            i = j + 1
            continue
        if t.type in ("fence", "code_block"):
            _add_code_block(doc, t.content)
            i += 1
            continue
        if t.type == "table_open":
            rows: list[list[str]] = []
            current_row: list[str] | None = None
            j = i + 1
            while j < len(tokens) and tokens[j].type != "table_close":
                tj = tokens[j]
                if tj.type == "tr_open":
                    current_row = []
                elif tj.type == "tr_close":
                    if current_row is not None:
                        rows.append(current_row); current_row = None
                elif tj.type == "inline" and current_row is not None:
                    text_parts: list[str] = []
                    for c in tj.children or []:
                        if c.type in ("text", "code_inline"):
                            text_parts.append(c.content)
                        elif c.type in ("softbreak", "hardbreak"):
                            text_parts.append(" ")
                    current_row.append("".join(text_parts).strip())
                j += 1
            _add_table(doc, rows)
            i = j + 1
            continue
        if t.type == "hr":
            doc.add_paragraph("─" * 40)
            i += 1
            continue
        i += 1

    doc.save(path)
